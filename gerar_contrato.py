import sys
import os
import traceback
from datetime import date

# Força o Python a usar UTF-8 para não dar erro com acentos no terminal
sys.stdout.reconfigure(encoding='utf-8')

# --- 0. FUNÇÃO DO LOGO ---
def exibir_logo():
    # Definição de Cores
    AZUL    = "\033[38;2;70;85;165m"
    AMARELO = "\033[38;2;235;192;70m"
    RESET   = "\033[0m"
    BLOCO   = "████████"
    
    print("\n")
    print(f"      {AZUL}{BLOCO}{RESET}      {AMARELO}{BLOCO}{RESET}")
    print(f"      {AZUL}{BLOCO}{RESET}      {AMARELO}{BLOCO}{RESET}")
    print(f"      {AZUL}{BLOCO}{RESET}      {AMARELO}{BLOCO}{RESET}")
    print("") 
    print(f"                      {AZUL}{BLOCO}{RESET}")
    print(f"                      {AZUL}{BLOCO}{RESET}")
    print(f"                      {AZUL}{BLOCO}{RESET}")
    print(f"          {AZUL}████{RESET}") 
    print(f"          {AZUL}████{RESET}")
    print(f"               {AZUL}{BLOCO}{RESET}")
    print(f"               {AZUL}{BLOCO}{RESET}")
    print(f"               {AZUL}{BLOCO}{RESET}")
    print("\n")
    print(f"{AZUL}    H o m o l o g{RESET}")
    print(f"{AMARELO}       S o l a r{RESET}")
    print("\n")
    print("--- INICIANDO PROCESSAMENTO ---")

# Chama o logo imediatamente
exibir_logo()

try:
    # Tenta importar as bibliotecas
    from openpyxl import load_workbook
    from docx import Document
    # --- NOVA BIBLIOTECA PARA PDF ---
    from docx2pdf import convert
except ImportError as e:
    print(f"\nERRO: Faltam bibliotecas!")
    print(f"Instale usando: pip install openpyxl python-docx docx2pdf")
    print(f"Detalhe: {e}")
    input("Pressione ENTER para sair...")
    sys.exit()

try:
    # --- 1. CONFIGURAÇÃO DE CAMINHOS ---
    
    # A) Onde estão os MODELOS e o EXE? (Pasta do Sistema)
    if getattr(sys, 'frozen', False):
        # Se estiver rodando como .exe, pega a pasta do executável
        pasta_sistema = os.path.dirname(sys.executable)
    else:
        # Se estiver rodando no VS Code/Terminal
        pasta_sistema = os.path.dirname(os.path.abspath(__file__))

    # B) Onde está a PLANILHA? (Vem do argumento do VBA)
    if len(sys.argv) > 1:
        caminho_excel = sys.argv[1]
    else:
        # Fallback para testes manuais
        caminho_excel = os.path.join(pasta_sistema, 'DADOS_DO_CLIENTE_version_1.xlsx')

    print(f"Sistema (EXE/Modelos) em: {pasta_sistema}")
    print(f"Lendo planilha em: {caminho_excel}")

    if not os.path.exists(caminho_excel):
        print(f"ERRO: O arquivo Excel não foi encontrado neste local: {caminho_excel}")
        print("Certifique-se de que o caminho enviado pelo VBA está correto.")
        raise FileNotFoundError

    # --- 2. CARREGAR EXCEL ---
    # Dica: Se der erro de permissão com colega, lembre-se de usar o truque do shutil.copy aqui
    wb = load_workbook(caminho_excel, data_only=True)
    sheet = wb.active

    hoje = date.today()
    meses_pt = ['Janeiro', 'Fevereiro', 'Março', 'Abril', 'Maio', 'Junho', 
                'Julho', 'Agosto', 'Setembro', 'Outubro', 'Novembro', 'Dezembro']
    nome_mes = meses_pt[hoje.month - 1]
    
    # --- 3. MAPEAMENTO DE DADOS ---
    dados = {
        "{{NOME}}": sheet['E6'].value,
        "{{CPF}}": sheet['E8'].value,
        "{{KWP}}": sheet['E5'].value,
        "{{ENDERECO}}": sheet['E11'].value,
        "{{CIDADE}}": sheet['E15'].value, 
        "{{CLASSIFICACAO}}": sheet['E21'].value,
        "{{CONTACONTRATO}}": sheet['E19'].value,
        "{{BAIRRO}}": sheet['E14'].value,
        "{{CEP}}": sheet['E13'].value,
        "{{DIA}}": hoje.day,
        "{{MES}}": nome_mes,  
        "{{ANO}}": hoje.year,
        "{{CONCESSIONARIA}}": sheet['E18'].value,
        "{{CPF_DO_REPRESENTANTE}}": sheet['I19'].value,
        "{{REPRESENTANTE}}": sheet['I18'].value,
    }
    
    # Limpeza: Transforma 'None' em "VAZIO" ou string limpa
    for chave, valor in dados.items():
        if valor is None:
            dados[chave] = "VAZIO"
        else:
            dados[chave] = str(valor).strip() # Remove espaços extras
        print(f"Lido > {chave}: {dados[chave]}")

    # --- 4. SELEÇÃO DO MODELO ---
    concessionaria = dados["{{CONCESSIONARIA}}"].upper()
    representante = dados["{{REPRESENTANTE}}"].upper()
    caminho_word_modelo = None 

    # IMPORTANTE: Usamos 'pasta_sistema' para procurar os modelos

    # Bloco CELPE
    if concessionaria == "CELPE 1" or concessionaria == "CELPE 2":
        if representante == "VAZIO":
            caminho_word_modelo = os.path.join(pasta_sistema, 'Procuração-celpe', 'MODELO-PROCURAÇÃO-Celpe-CPF.docx')
        else:
            caminho_word_modelo = os.path.join(pasta_sistema, 'Procuração-celpe', 'MODELO-PROCURAÇÃO-Celpe-CNPJ.docx')

    # Bloco COELBA
    elif concessionaria == "COELBA":
        if representante == "VAZIO":
            caminho_word_modelo = os.path.join(pasta_sistema, 'Procuração-Coelba', 'MODELO-PROCURAÇÃO-Coelba-CPF.docx')
        else:
            caminho_word_modelo = os.path.join(pasta_sistema, 'Procuração-Coelba', 'MODELO-PROCURAÇÃO-Coelba-CNPJ.docx')

    # Bloco COSERN
    elif concessionaria == "COSERN":
        if representante == "VAZIO":
            caminho_word_modelo = os.path.join(pasta_sistema, 'Procuração-Cosern', 'MODELO-PROCURAÇÃO-Cosern-CPF.docx')
        else:
            caminho_word_modelo = os.path.join(pasta_sistema, 'Procuração-Cosern', 'MODELO-PROCURAÇÃO-Cosern-CNPJ.docx')

    # Bloco EQUATORIAL
    elif concessionaria == "EQUATORIAL":
        if representante == "VAZIO":
            caminho_word_modelo = os.path.join(pasta_sistema, 'Procuração-Equatorial', 'MODELO-PROCURAÇÃO-Equatorial-CPF.docx')
        else:
            caminho_word_modelo = os.path.join(pasta_sistema, 'Procuração-Equatorial', 'MODELO-PROCURAÇÃO-Equatorial-CNPJ.docx')

    # --- 5. VERIFICAÇÃO E GERAÇÃO ---
    if caminho_word_modelo and os.path.exists(caminho_word_modelo):
        print(f"Modelo encontrado: {caminho_word_modelo}")
        doc = Document(caminho_word_modelo)
        
        # Substituição em parágrafos
        for paragrafo in doc.paragraphs:
            for codigo, valor in dados.items():
                if codigo in paragrafo.text:
                    paragrafo.text = paragrafo.text.replace(codigo, str(valor))
        
        # Substituição em tabelas
        for tabela in doc.tables:
            for linha in tabela.rows:
                for celula in linha.cells:
                    for codigo, valor in dados.items():
                        if codigo in celula.text:
                            celula.text = celula.text.replace(codigo, str(valor))

        # --- SALVAR O ARQUIVO WORD ---
        pasta_saida = os.path.dirname(caminho_excel)
        nome_cliente = dados["{{NOME}}"] if dados["{{NOME}}"] != "VAZIO" else "Cliente"
        nome_cliente_limpo = "".join(x for x in nome_cliente if x.isalnum() or x in " _-")
        
        # Define os nomes dos dois arquivos
        nome_arquivo_docx = os.path.join(pasta_saida, f"Procuracao_{concessionaria}_{nome_cliente_limpo}.docx")
        nome_arquivo_pdf = os.path.join(pasta_saida, f"Procuracao_{concessionaria}_{nome_cliente_limpo}.pdf")
        
        # Salva o Word
        doc.save(nome_arquivo_docx)
        print(f"\nSUCESSO! Word salvo em:\n{nome_arquivo_docx}")
        
        # --- 6. GERAR O PDF ---
        print("\nGerando PDF (Aguarde o Word processar)...")
        try:
            # Converte o arquivo DOCX que acabamos de criar
            convert(nome_arquivo_docx, nome_arquivo_pdf)
            print(f"SUCESSO! PDF salvo em:\n{nome_arquivo_pdf}")
        except Exception as e_pdf:
            print(f"ERRO AO GERAR PDF: {e_pdf}")
            print("Dica: Verifique se o Microsoft Word está instalado e feche janelas de diálogo abertas.")

    else:
        print("\n[!] ATENÇÃO: Não foi possível gerar os arquivos.")
        if caminho_word_modelo is None:
            print(f"Motivo: Concessionária '{concessionaria}' não configurada ou desconhecida.")
        else:
            print(f"Motivo: O arquivo de modelo não existe no caminho:\n{caminho_word_modelo}")

except Exception:
    print("\nXXX OCORREU UM ERRO GRAVE XXX")
    traceback.print_exc()

print("\n--------------------------------")
input("Pressione ENTER para fechar...")